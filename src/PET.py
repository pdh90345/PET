import streamlit as st
import openai
import chardet
import os
import base64
import PyPDF2
from io import BytesIO
from docx import Document
from MySQL_Connector import connect_to_db

st.title("P.E.T 문제 생성기")
st.write(
    "텍스트 파일을 업로드하고 해당 내용을 기반으로 객관식, 참/거짓 또는 서술형의 문제를 생성합니다."
)

input_id = st.text_input("사용자의 ID를 입력하세요: ")

connection = connect_to_db()
cursor = connection.cursor(dictionary=True)

if "id" not in st.session_state:
    st.session_state.id = ""
st.session_state.id = input_id
button = st.button("아이디 확인")


def check_login(input_id):
    if input_id is "":
        return 0

    sql = "SELECT user_id FROM membership WHERE user_id = %s and login_check = 1;"
    cursor.execute(sql, (input_id,))
    result = cursor.fetchall()
    if result:
        st.write("아이디가 확인 되었습니다.")
    else:
        st.write("로그인한 아이디가 아닙니다.")
    return result


def get_generated_questions():
    return st.session_state["questions"]


def generate_questions_gpt35_turbo(text, num_questions, question_type, num_options=4):
    if question_type == "객관식":
        gpt_prompt = f"다음과 같이 {num_questions}개의 객관식 문제를 생성해주세요. 각 문제에는 {num_options}개의 선택지가 있으며, 하나의 선택지만 올바른 답입니다. 문제의 답은 반드시 중간에 출력되면 안되고 문제의 답은 무조건 마지막 질문 생성 후에 'Ans:'뒤에 모아서 한번에 출력해주세요."
        role_content = "You are a helpful assistant that generates multiple choice questions based on the provided text."
    elif question_type == "참/거짓형":
        gpt_prompt = f"다음과 같이 {num_questions}개의 참/거짓형 문제를 생성해주세요. 가능한 답은 'T' 또는 'F'로 구성하며, 문제의 답은 반드시 중간에 보여지면 안되고 문제의 답은 무조건 마지막 질문 생성 후에 'Ans:'뒤에 모아서 한번에 출력해주세요."
        role_content = "You are a helpful assistant that generates true or false questions based on the provided text."
    else:
        gpt_prompt = f"다음과 같이 {num_questions}개의 서술형 문제를 생성해주세요. 문제의 답은 질문에 맞는 간단한 문장이며, 문제의 답은 반드시 중간에 보여지면 안되고 문제의 답은 무조건 마지막 질문 생성 후에 'Ans:'뒤에 모아서 한번에 출력해주세요."
        role_content = "You are a helpful assistant that generates open-ended questions based on the provided text."

    prompt = [
        {"role": "system", "content": role_content},
        {"role": "user", "content": text},
        {"role": "assistant", "content": gpt_prompt},
    ]

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=prompt,
        temperature=0.5,
        max_tokens=2000,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0,
    )

    cursor = connection.cursor(dictionary=True)

    query = "UPDATE membership SET question_current = question_current + %s WHERE user_id = %s"

    cursor.execute(
        query,
        (
            num_questions,
            input_id,
        ),
    )

    connection.commit()

    answer = response["choices"][0]["message"]["content"].strip()

    return answer


def read_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


def read_docx(file):
    doc = Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


def text_to_docx(questions):
    doc = Document()
    doc.add_paragraph(questions)
    byte_io = BytesIO()
    doc.save(byte_io)
    return byte_io.getvalue()


def get_binary_file_downloader_html(bin_file, file_label="File"):
    with st.spinner("파일 생성 중..."):
        b64 = base64.b64encode(bin_file).decode()
        href = f'<a href="data:application/octet-stream;base64,{b64}" download="{file_label}.docx">파일 다운로드 .docx</a>'
    return href


def main():
    if check_login(input_id):
        uploaded_file = st.file_uploader(
            "텍스트 파일을 업로드하세요.", type=["txt", "pdf", "docx"]
        )
        st.warning(
            "텍스트가 한글 2000자 이상이라면 문제가 생성되지 않을 수 있습니다. 텍스트를 줄여주세요.",
            icon="⚠️",
        )
        if uploaded_file:
            file_extension = uploaded_file.name.split(".")[-1].lower()

            if file_extension == "txt":
                raw_data = uploaded_file.read()
                detected_encoding = chardet.detect(raw_data)
                text = raw_data.decode(detected_encoding["encoding"])
            elif file_extension == "pdf":
                text = read_pdf(uploaded_file)
            elif file_extension == "docx":
                text = read_docx(uploaded_file)
            else:
                st.error(
                    "지원되지 않는 파일 형식입니다. TXT, PDF 또는 DOCX 파일을 업로드해주세요."
                )

            question_type = st.selectbox(
                "생성할 문제 유형을 선택하세요.", ["객관식", "참/거짓형", "서술형"]
            )

            if question_type == "객관식":
                num_options = st.number_input(
                    "각 문제에 대한 선택지 수를 선택하세요 :",
                    min_value=2,
                    max_value=10,
                    value=4,
                    step=1,
                )
            else:
                num_options = 4

            num_questions = st.number_input(
                "생성할 문제 수를 입력하세요 :", min_value=1, value=5
            )

            # # API key 입력
            # api_key = st.text_input("OpenAI API 키를 입력하세요 :", type="password")
            # if not api_key:
            #     st.warning("계속하기 전에 API 키를 입력해주세요. \n\n https://platform.openai.com/account/api-keys 에서 찾을 수 있습니다.")
            #     st.stop()
            api_key = ""

            openai.api_key = api_key

            if "questions" not in st.session_state:
                st.session_state["questions"] = ""

            if st.button("문제 생성"):
                st.session_state["questions"] = generate_questions_gpt35_turbo(
                    text, num_questions, question_type, num_options
                )
                question_parts = (
                    st.session_state["questions"].replace("*", "\\*").split("Ans:")
                )
                st.session_state["Q"] = question_parts[0]
                st.session_state["A"] = question_parts[1]
                st.success("문제가 생성되었습니다.")

            if st.button("문제 출력"):
                if "Q" in st.session_state and "A" in st.session_state:
                    st.write("질문:\n", st.session_state["Q"])
                    st.write("답:", st.session_state["A"])

            if st.button("질문 저장"):
                if st.session_state["Q"]:
                    docx_file = text_to_docx(st.session_state["Q"])
                    st.markdown(
                        get_binary_file_downloader_html(docx_file, "질문"),
                        unsafe_allow_html=True,
                    )
                else:
                    st.write(
                        "문제가 생성되지 않았습니다. 먼저 '문제 생성' 버튼을 눌러 질문을 생성하세요."
                    )

            if st.button("답 저장"):
                if st.session_state["A"]:
                    docx_file = text_to_docx(st.session_state["A"])
                    st.markdown(
                        get_binary_file_downloader_html(docx_file, "답"),
                        unsafe_allow_html=True,
                    )
                else:
                    st.write(
                        "문제가 생성되지 않았습니다. 먼저 '문제 생성' 버튼을 눌러 질문을 생성하세요."
                    )
            return get_generated_questions()


if __name__ == "__main__":
    main()
